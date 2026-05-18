"""
report_generator.py

Generates an ISO 14971-style Risk Management Report (DOCX) from a
completed RMF risk table, device information, RMP configuration, and
optional follow-up Q&A answers.

Public API
----------
    generate_rmf_report(df, device_info, rmp_config, followup_qa) -> bytes

The returned bytes can be passed directly to st.download_button.
No filesystem writes are performed; everything is in-memory.

RAG integration points are marked with  # RAG HOOK  comments so that
future context injection (standards passages, retrieved cases, etc.)
can be located and added without restructuring the module.
"""
from __future__ import annotations

import io
from collections import Counter
from datetime import datetime
from typing import Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# ── Low-level DOCX helpers ─────────────────────────────────────────────────────

def _bold_cell(cell) -> None:
    """Make every run in a table cell bold (used for header rows)."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _add_kv(doc: Document, key: str, value: str) -> None:
    """Emit a  Key: value  paragraph where the key is bold."""
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(str(value) if value else "N/A")


def _add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(str(text), style="List Bullet")


def _add_text(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _add_table(doc: Document, headers: list[str], rows: list[list]) -> None:
    """Render a grid table with a bold header row."""
    if not headers:
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _bold_cell(hdr_cells[i])

    for ri, row_data in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row_data):
            # Safely convert NaN / None to empty string
            if val is None:
                cells[ci].text = ""
            elif isinstance(val, float) and val != val:  # NaN check
                cells[ci].text = ""
            else:
                cells[ci].text = str(val)


# ── Column-name normalisation ──────────────────────────────────────────────────
# The DataFrame may arrive with Title Case columns (from LLM output) or
# snake_case columns (from the database).  _col_map() resolves whichever
# form is present so all downstream code uses display names as keys.

_DISPLAY_TO_SNAKE: dict[str, str] = {
    "Hazard":               "hazard",
    "Hazardous Situation":  "hazardous_situation",
    "Possible Harm":        "possible_harm",
    "Severity":             "severity",
    "Probability":          "probability",
    "Initial Risk Level":   "initial_risk_level",
    "Risk Control Measure": "risk_control_measure",
    "Residual Risk":        "residual_risk",
    "Verification Method":  "verification_method",
    "Status":               "status",
}


def _col_map(df: pd.DataFrame) -> dict[str, str]:
    """Return {display_name: actual_column_name} for every column present in df."""
    existing = set(df.columns)
    result: dict[str, str] = {}
    for display, snake in _DISPLAY_TO_SNAKE.items():
        if display in existing:
            result[display] = display
        elif snake in existing:
            result[display] = snake
    return result


# ── Statistics ─────────────────────────────────────────────────────────────────

def _compute_stats(df: pd.DataFrame) -> dict:
    """Compute summary counts used throughout the report."""
    if df.empty:
        return {
            "total": 0,
            "hazards": [],
            "controls": [],
            "init_high": 0,
            "init_med":  0,
            "init_low":  0,
            "res_high":  0,
            "res_med":   0,
            "res_low":   0,
            "unacceptable": 0,
        }

    c = _col_map(df)

    def _normalised_series(display: str) -> pd.Series:
        col = c.get(display)
        if not col:
            return pd.Series([], dtype=str)
        return df[col].fillna("").astype(str).str.strip().str.title()

    init_counts = Counter(_normalised_series("Initial Risk Level"))
    res_counts  = Counter(_normalised_series("Residual Risk"))

    hazard_col  = c.get("Hazard")
    control_col = c.get("Risk Control Measure")

    return {
        "total":   len(df),
        "hazards": df[hazard_col].fillna("").drop_duplicates().tolist() if hazard_col else [],
        "controls": df[control_col].fillna("").drop_duplicates().tolist() if control_col else [],
        "init_high": init_counts.get("High",   0),
        "init_med":  init_counts.get("Medium", 0),
        "init_low":  init_counts.get("Low",    0),
        "res_high":  res_counts.get("High",    0),
        "res_med":   res_counts.get("Medium",  0),
        "res_low":   res_counts.get("Low",     0),
        "unacceptable": res_counts.get("High", 0),
    }


# ── Dynamic text generators ────────────────────────────────────────────────────

def _text_risk_analysis_summary(
    device_name: str,
    device_type: str,
    stats: dict,
    followup_qa: list,
) -> str:
    qa_note = (
        " The analysis incorporated device-specific information collected via a"
        " structured follow-up Q&A session covering operating environment, user"
        " population, critical failure modes, and applicable safety requirements."
        if any(item.get("answer", "").strip() for item in followup_qa)
        else ""
    )
    return (
        f"Risk analysis for the {device_name} ({device_type}) identified"
        f" {stats['total']} risk record(s) from {len(stats['hazards'])} unique"
        f" hazard source(s). {stats['init_high']} hazard(s) were initially assessed"
        f" as high risk.{qa_note} The analysis followed the ISO 14971 framework for"
        " hazard identification, risk estimation, and risk evaluation."
    )


def _text_benefit_risk(
    device_name: str,
    intended_use: str,
    stats: dict,
) -> str:
    total = stats["total"]
    if total == 0:
        return (
            f"A benefit-risk assessment was conducted for the {device_name}. "
            "Based on qualitative assessment, the intended clinical benefits outweigh "
            "the residual risks associated with the device."
        )
    reduction = ""
    if stats["init_high"] > 0 and stats["res_high"] < stats["init_high"]:
        reduction = (
            f" Risk control measures reduced high-risk items from"
            f" {stats['init_high']} to {stats['res_high']}, demonstrating the"
            " effectiveness of the applied controls."
        )
    elif stats["res_high"] == 0:
        reduction = " All identified risks have been reduced to acceptable levels."
    return (
        f"The {device_name} is intended for {intended_use}. "
        f"A total of {total} hazard(s) were identified and analysed."
        f" {stats['init_high']} were initially classified as high risk.{reduction} "
        "The clinical benefits of the device — including its diagnostic or therapeutic"
        " value for the intended patient population — are assessed to outweigh the"
        " residual risks when the device is used in accordance with its intended use"
        " and the accompanying instructions for use."
    )


def _text_residual_risk(stats: dict, rmp_config: dict) -> str:
    if stats["total"] == 0:
        return "No risk records were available to evaluate overall residual risk."

    method = rmp_config.get("residual_risk_method", "").strip()
    basis  = rmp_config.get("residual_risk_basis", [])

    if stats["res_high"] > 0:
        body = (
            f"Overall residual risk evaluation identified {stats['res_high']} item(s)"
            " with residual risk classified as High. These items require additional"
            " documented justification before the device can be declared safe for its"
            " intended use."
        )
    else:
        body = (
            f"After applying all risk control measures, residual risks are distributed"
            f" as follows: {stats['res_high']} High, {stats['res_med']} Medium,"
            f" {stats['res_low']} Low."
            " The overall residual risk is assessed as acceptable in accordance with"
            " the defined risk acceptability criteria."
        )

    if method:
        body += f"\n\nEvaluation Method: {method}"
    if basis:
        body += f"\n\nEvaluation Basis: {', '.join(basis)}."

    return body


# ── Section builders ───────────────────────────────────────────────────────────

def _section_title(
    doc: Document,
    device_info: dict,
    rmp_config: dict,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Risk Management Report")
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph()
    _add_kv(doc, "Device Name",  device_info.get("device_name",  ""))
    _add_kv(doc, "Device Type",  device_info.get("device_type",  ""))
    _add_kv(doc, "Intended Use", device_info.get("intended_use", ""))
    _add_kv(doc, "Report Date",  datetime.now().strftime("%B %d, %Y"))
    _add_kv(doc, "Standard",     "ISO 14971:2019 — Medical devices — Application of risk management")
    _add_kv(doc, "Document Status", "Draft — Requires Expert Review Before Regulatory Use")

    team = rmp_config.get("team_members", "").strip()
    if team:
        doc.add_paragraph()
        doc.add_heading("Prepared By", level=3)
        for line in team.splitlines():
            line = line.strip().lstrip("-").strip()
            if line:
                _add_bullet(doc, line)

    doc.add_page_break()


def _section_device_description(
    doc: Document,
    device_info: dict,
    rmp_config: dict,
) -> None:
    doc.add_heading("1. Device Description", level=1)
    _add_kv(doc, "Device Name",  device_info.get("device_name",  ""))
    _add_kv(doc, "Device Type",  device_info.get("device_type",  ""))
    _add_kv(doc, "Intended Use", device_info.get("intended_use", ""))

    lifecycle = rmp_config.get("lifecycle_scope", [])
    if lifecycle:
        doc.add_paragraph()
        _add_text(doc, "Lifecycle Phases Covered by This Risk Management Plan:", bold=True)
        for phase in lifecycle:
            _add_bullet(doc, phase)


def _section_processes_and_documents(
    doc: Document,
    rmp_config: dict,
) -> None:
    doc.add_heading("2. Relevant Processes and Documents", level=1)
    _add_text(doc, "This Risk Management Report was prepared in accordance with:")
    for ref in [
        "ISO 14971:2019 — Application of risk management to medical devices",
        "IEC 62304:2006+AMD1:2015 — Medical device software — Software life cycle processes (where applicable)",
        "IEC 62366-1:2015 — Medical devices — Usability engineering (where applicable)",
    ]:
        _add_bullet(doc, ref)

    methods = rmp_config.get("verification_methods", [])
    if methods:
        doc.add_paragraph()
        _add_text(doc, "Verification Method Library:", bold=True)
        for m in methods:
            _add_bullet(doc, m)

    criteria = rmp_config.get("risk_acceptability_criteria", {})
    if criteria:
        doc.add_paragraph()
        _add_text(doc, "Risk Acceptability Criteria:", bold=True)
        if criteria.get("type") == "custom":
            doc.add_paragraph(criteria.get("content", ""))
        else:
            doc.add_paragraph(
                "Risk acceptability is defined using a probability × severity matrix. "
                "Cells in the matrix are assigned LOW, MEDIUM, or HIGH risk levels "
                "according to the configured Risk Management Plan."
            )

    # RAG HOOK: applicable standards sections could be retrieved and inserted here


def _section_risk_analysis(
    doc: Document,
    df: pd.DataFrame,
    device_info: dict,
    stats: dict,
    followup_qa: list,
) -> None:
    doc.add_heading("3. Risk Analysis", level=1)
    doc.add_paragraph(_text_risk_analysis_summary(
        device_info.get("device_name", ""),
        device_info.get("device_type", ""),
        stats,
        followup_qa,
    ))

    # 3.1 Preliminary Hazards Analysis
    doc.add_heading("3.1 Preliminary Hazards Analysis", level=2)
    # RAG HOOK: retrieved ISO 14971 hazard taxonomy could seed this list
    hazards = [str(h) for h in stats.get("hazards", []) if str(h).strip()]
    if hazards:
        doc.add_paragraph(f"{len(hazards)} unique hazard source(s) identified:")
        for h in hazards:
            _add_bullet(doc, h)
    else:
        doc.add_paragraph("No hazard data available.")

    # 3.2 Failure Modes
    doc.add_heading("3.2 Failure Modes", level=2)
    if not df.empty:
        c = _col_map(df)
        haz_col  = c.get("Hazard")
        sit_col  = c.get("Hazardous Situation")
        harm_col = c.get("Possible Harm")
        if haz_col and sit_col and harm_col:
            subset = (
                df[[haz_col, sit_col, harm_col]]
                .fillna("")
                .drop_duplicates()
            )
            _add_table(
                doc,
                ["Hazard", "Hazardous Situation", "Possible Harm"],
                subset.values.tolist(),
            )
        else:
            doc.add_paragraph("Failure mode columns not available in the risk table.")
    else:
        doc.add_paragraph("No risk records available.")

    # 3.3 FMEA Summary
    doc.add_heading("3.3 FMEA Summary", level=2)
    _add_table(
        doc,
        ["Risk Level", "Initial Count", "Residual Count"],
        [
            ["High",   stats["init_high"], stats["res_high"]],
            ["Medium", stats["init_med"],  stats["res_med"]],
            ["Low",    stats["init_low"],  stats["res_low"]],
            ["Total",  stats["total"],     stats["total"]],
        ],
    )

    # 3.4 Device-Specific Information (from follow-up Q&A)
    # RAG HOOK: retrieved risk cases could be appended here as additional context
    non_empty_qa = [
        (item.get("question", ""), item.get("answer", ""))
        for item in followup_qa
        if item.get("answer", "").strip()
    ]
    if non_empty_qa:
        doc.add_heading("3.4 Device-Specific Risk Information", level=2)
        doc.add_paragraph(
            "The following device-specific information was collected to inform the risk analysis:"
        )
        for q, a in non_empty_qa:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(q + " ").bold = True
            p.add_run(a)


def _section_risk_control_measures(
    doc: Document,
    df: pd.DataFrame,
    stats: dict,
) -> None:
    doc.add_heading("4. Risk Control Measures", level=1)
    doc.add_paragraph(
        f"{len(stats['controls'])} distinct risk control measure(s) were applied "
        "to reduce identified risks to acceptable levels."
    )

    if not df.empty:
        c = _col_map(df)
        haz_col  = c.get("Hazard")
        ctrl_col = c.get("Risk Control Measure")
        vrfy_col = c.get("Verification Method")
        res_col  = c.get("Residual Risk")
        if haz_col and ctrl_col and vrfy_col and res_col:
            subset = (
                df[[haz_col, ctrl_col, vrfy_col, res_col]]
                .fillna("")
                .drop_duplicates()
            )
            _add_table(
                doc,
                ["Hazard", "Risk Control Measure", "Verification Method", "Residual Risk"],
                subset.values.tolist(),
            )
    # RAG HOOK: standards-based control rationale could be injected per row


def _section_risk_matrix_summary(
    doc: Document,
    stats: dict,
) -> None:
    doc.add_heading("5. Risk Matrix Summary", level=1)
    total = stats["total"]
    if total == 0:
        doc.add_paragraph("No risk records available.")
        return

    def _pct(n: int) -> str:
        return f"{n} ({round(n / total * 100)}%)"

    doc.add_paragraph(
        f"Distribution of risk levels across all {total} identified risk record(s)."
    )
    _add_table(
        doc,
        ["Risk Level", "Initial Risk", "Residual Risk"],
        [
            ["High",   _pct(stats["init_high"]), _pct(stats["res_high"])],
            ["Medium", _pct(stats["init_med"]),  _pct(stats["res_med"])],
            ["Low",    _pct(stats["init_low"]),  _pct(stats["res_low"])],
        ],
    )


def _section_unacceptable_risks(
    doc: Document,
    df: pd.DataFrame,
    stats: dict,
) -> None:
    doc.add_heading("6. Summary of Risks and Unacceptable Risks", level=1)
    doc.add_paragraph(f"Total risk records: {stats['total']}")
    doc.add_paragraph(
        f"Residual high-risk items (require further action): {stats['unacceptable']}"
    )

    if stats["unacceptable"] > 0 and not df.empty:
        c = _col_map(df)
        res_col  = c.get("Residual Risk")
        haz_col  = c.get("Hazard")
        harm_col = c.get("Possible Harm")
        if res_col and haz_col:
            mask = (
                df[res_col].fillna("").astype(str).str.strip().str.title() == "High"
            )
            high_df = df[mask]
            if not high_df.empty:
                _add_text(doc, "Items requiring further documented justification:", bold=True)
                show_cols = [col for col in [haz_col, harm_col] if col]
                for _, row in high_df.iterrows():
                    _add_bullet(doc, " — ".join(str(row[col]) for col in show_cols))
    elif stats["unacceptable"] == 0 and stats["total"] > 0:
        doc.add_paragraph(
            "All risk records have been reduced to acceptable residual risk levels "
            "through the application of risk control measures."
        )


def _section_benefit_risk(
    doc: Document,
    device_info: dict,
    stats: dict,
) -> None:
    doc.add_heading("7. Benefit-Risk Assessment", level=1)
    # RAG HOOK: retrieved clinical evidence or benefit statements could be prepended
    doc.add_paragraph(_text_benefit_risk(
        device_info.get("device_name",  ""),
        device_info.get("intended_use", ""),
        stats,
    ))


def _section_overall_residual_risk(
    doc: Document,
    stats: dict,
    rmp_config: dict,
) -> None:
    doc.add_heading("8. Overall Residual Risk", level=1)
    # RAG HOOK: standards-based residual risk acceptance criteria could be cited here
    doc.add_paragraph(_text_residual_risk(stats, rmp_config))


def _section_appendix(
    doc: Document,
    df: pd.DataFrame,
) -> None:
    doc.add_heading("Appendix — Full Risk Table", level=1)
    if df.empty:
        doc.add_paragraph("No risk records available.")
        return

    c = _col_map(df)
    ordered_display = [
        "Hazard", "Hazardous Situation", "Possible Harm",
        "Severity", "Probability", "Initial Risk Level",
        "Risk Control Measure", "Residual Risk",
        "Verification Method", "Status",
    ]
    actual_headers = [h for h in ordered_display if h in c]
    actual_cols    = [c[h] for h in actual_headers]

    if not actual_cols:
        doc.add_paragraph("Risk table data unavailable.")
        return

    rows = df[actual_cols].fillna("").values.tolist()
    _add_table(doc, actual_headers, rows)


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_rmf_report(
    df: pd.DataFrame,
    device_info: dict,
    rmp_config: dict,
    followup_qa: Optional[list] = None,
) -> bytes:
    """
    Build an ISO 14971-style Risk Management Report and return it as bytes.

    Parameters
    ----------
    df : pd.DataFrame
        The RMF risk table.  Accepts Title Case column names (from LLM output)
        or snake_case column names (from the database).
    device_info : dict
        Must contain keys: device_name, intended_use, device_type.
    rmp_config : dict
        The saved RMP configuration (result of get_latest_rmp_config()).
        Pass an empty dict {} if no configuration is available.
    followup_qa : list[dict], optional
        [{"question": "...", "answer": "..."}, ...] from the RMF wizard.
        Pass None or [] if not available.

    Returns
    -------
    bytes
        Raw DOCX file content.  Pass directly to st.download_button(data=...).
    """
    if followup_qa is None:
        followup_qa = []
    if rmp_config is None:
        rmp_config = {}

    stats = _compute_stats(df)
    doc   = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    _section_title(doc, device_info, rmp_config)
    _section_device_description(doc, device_info, rmp_config)
    doc.add_paragraph()
    _section_processes_and_documents(doc, rmp_config)
    doc.add_paragraph()
    _section_risk_analysis(doc, df, device_info, stats, followup_qa)
    doc.add_paragraph()
    _section_risk_control_measures(doc, df, stats)
    doc.add_paragraph()
    _section_risk_matrix_summary(doc, stats)
    doc.add_paragraph()
    _section_unacceptable_risks(doc, df, stats)
    doc.add_paragraph()
    _section_benefit_risk(doc, device_info, stats)
    doc.add_paragraph()
    _section_overall_residual_risk(doc, stats, rmp_config)
    doc.add_page_break()
    _section_appendix(doc, df)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
